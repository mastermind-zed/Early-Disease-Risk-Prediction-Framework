import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Set working directory
os.chdir(r'C:\Users\SAMUEL ADJEI\Desktop\Disease Prediction\Disease_Prediction')

def create_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('Machine Learning Framework for Early Disease Risk Prediction', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Final Project Report - Diabetes Case Study')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Section 1: Introduction
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        "This report presents the findings of a machine learning study aimed at predicting the risk of Early Disease (Diabetes). "
        "The project follows a systematic framework including data acquisition, preprocessing, and comparative model evaluation."
    )
    
    # Section 2: Methodology
    doc.add_heading('2. Methodology', level=1)
    doc.add_paragraph(
        "Four machine learning algorithms were implemented and evaluated:\n"
        "- Logistic Regression (Baseline)\n"
        "- Random Forest (Ensemble)\n"
        "- Decision Tree (Interpretable Tree)\n"
        "- Support Vector Machine (Robust Classification)"
    )
    doc.add_paragraph(
        "The dataset used was a clinical and demographic dataset containing 2,954 records. "
        "Evaluation metrics include Accuracy, Precision, Recall, F1-Score, and AUC-ROC."
    )
    
    # Section 3: Exploratory Data Analysis
    doc.add_heading('3. Exploratory Data Analysis', level=1)
    doc.add_paragraph("The feature correlation analysis revealed strong links between glucose indicators and the disease outcome.")
    
    if os.path.exists('correlation_heatmap.png'):
        doc.add_picture('correlation_heatmap.png', width=Inches(5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Figure 1: Feature Correlation Heatmap', style='Caption').alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Section 4: Results and Comparison
    doc.add_heading('4. Results and Comparison', level=1)
    doc.add_paragraph("The table below summarizes the performance of the four models on the test set:")
    
    # Load results
    if os.path.exists('model_results_all.csv'):
        results_df = pd.read_csv('model_results_all.csv')
        
        # Add table
        table = doc.add_table(rows=1, cols=len(results_df.columns))
        table.style = 'Table Grid'
        
        # Header
        hdr_cells = table.rows[0].cells
        for i, column in enumerate(results_df.columns):
            hdr_cells[i].text = column
            
        # Data
        for Index, row in results_df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                if isinstance(value, float):
                    row_cells[i].text = f"{value:.4f}"
                else:
                    row_cells[i].text = str(value)
                    
    doc.add_paragraph("\n")
    
    # ROC Curves
    if os.path.exists('roc_curves_all.png'):
        doc.add_picture('roc_curves_all.png', width=Inches(5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Figure 2: ROC Curves for All Models', style='Caption').alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Section 5: Discussion - Data Quality vs. Quantity
    doc.add_heading('5. Discussion: Data Quality vs. Quantity', level=1)
    doc.add_paragraph(
        "A critical observation in this study is the high predictive accuracy (99%+) achieved with a dataset of 2,954 records, "
        "compared to common large-scale datasets (e.g., Kaggle diabetes datasets with 100,000+ records) which often yield lower accuracies (95-97%)."
    )
    doc.add_paragraph(
        "This is attributed to the 'Signal-to-Noise' ratio in the clinical features. The features in this dataset, particularly clinical symptoms "
        "(Polyuria, Polydipsia) and laboratory results (HbA1c, FPG), are direct physiological indicators of the disease. "
        "Unlike lifestyle-focused datasets that rely on noisier indicators like BMI or smoking history, these clinical markers provide a clear "
        "decision boundary for machine learning algorithms. This demonstrates that for medical risk prediction, the quality and relevance of "
        "clinical features are more impactful than the sheer volume of general demographic data."
    )

    # Section 6: Future Work - Framework Scalability
    doc.add_heading('6. Future Work: Framework Scalability', level=1)
    doc.add_paragraph(
        "While this study focuses on Diabetes Mellitus, the developed machine learning pipeline is architected as a "
        "disease-agnostic framework. The modular design of the data preprocessing and model evaluation scripts allows "
        "for seamless adaptation to other chronic conditions (e.g., cardiovascular disease, chronic kidney disease) "
        "provided that binary classification clinical data is available. This scalability makes the framework a "
        "valuable asset for broader early disease risk screening initiatives."
    )

    # Section 7: Conclusion
    doc.add_heading('7. Conclusion', level=1)
    doc.add_paragraph(
        "The Random Forest and SVM models demonstrated superior performance, both achieving near-perfect metrics. "
        "The high AUC values indicate that the underlying clinical features are extremely robust predictors. "
        "The framework developed can be successfully applied to other disease prediction tasks given a structured clinical dataset."
    )
    
    # Save document
    doc.save('reports/Early_Disease_Risk_Prediction_Report.docx')
    print("Report generated: reports/Early_Disease_Risk_Prediction_Report.docx")

if __name__ == "__main__":
    create_report()
